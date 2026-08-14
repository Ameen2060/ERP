"""Transaction document attachments: secure storage, a lightweight template-free document
intelligence pass (PDF / Excel / CSV / Word text → fields), amount matching against the
owning transaction, and a full audit trail.

Files live on disk under settings.attachments_dir/<entity_type>/<entity_id>/<id><ext>; the DB
row is the source of truth. Deletes are soft (is_deleted) so nothing leaves the audit trail.
"""

from __future__ import annotations

import hashlib
import io
import json
import os
import re
from datetime import datetime, timezone
from decimal import Decimal, InvalidOperation

from sqlalchemy import select
from sqlalchemy.orm import Session

from ..config import get_settings
from . import storage
from ..models import (
    Attachment,
    AttachmentEvent,
    BankStatementLine,
    BillPayment,
    CustomerAdvance,
    CustomerCreditNote,
    CustomerPayment,
    EInvoice,
    Expense,
    FixedAsset,
    JournalEntry,
    PayrollRun,
    Project,
    SalesInvoice,
    StockMovement,
    VendorAdvance,
    VendorBill,
    VendorCreditNote,
)


class AttachmentError(ValueError):
    """Raised on a rejected upload (bad type/size), missing entity, or bad operation."""


# ── Allow-list & MIME ─────────────────────────────────────────────────────────────────────
ALLOWED_EXT = {"pdf", "jpg", "jpeg", "png", "gif", "webp", "xlsx", "xls", "csv", "docx", "doc", "txt"}
_MIME = {
    "pdf": "application/pdf", "jpg": "image/jpeg", "jpeg": "image/jpeg", "png": "image/png",
    "gif": "image/gif", "webp": "image/webp", "csv": "text/csv", "txt": "text/plain",
    "xlsx": "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
    "xls": "application/vnd.ms-excel",
    "docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    "doc": "application/msword",
}
_IMAGE_EXT = {"jpg", "jpeg", "png", "gif", "webp"}
_ENTITY_TYPES = {"journal_entry", "sales_invoice", "vendor_bill", "customer_payment",
                 "bill_payment", "bank_statement_line", "fixed_asset", "stock_movement",
                 "payroll_run", "expense", "customer_credit_note", "vendor_credit_note", "project",
                 "customer_advance", "vendor_advance", "einvoice"}


def _now() -> datetime:
    return datetime.now(timezone.utc)


def _iso(dt: datetime | None) -> str | None:
    return dt.isoformat() if dt else None


def _ext_of(name: str) -> str:
    return (os.path.splitext(name)[1].lstrip(".") or "").lower()


# ── Owning-transaction lookup + expected amount ──────────────────────────────────────────────
def _entity_amount(db: Session, entity_type: str, entity_id: str) -> Decimal | None:
    """The transaction's own total, used to compare against an extracted invoice total."""
    if entity_type == "journal_entry":
        je = db.get(JournalEntry, entity_id)
        return Decimal(sum((ln.debit for ln in je.lines), Decimal(0))) if je else None
    if entity_type == "sales_invoice":
        inv = db.get(SalesInvoice, entity_id)
        return Decimal(inv.grand_total) if inv else None
    if entity_type == "vendor_bill":
        bill = db.get(VendorBill, entity_id)
        return Decimal(bill.grand_total) if bill else None
    if entity_type == "customer_payment":
        p = db.get(CustomerPayment, entity_id)
        return Decimal(p.amount) if p else None
    if entity_type == "bill_payment":
        p = db.get(BillPayment, entity_id)
        return Decimal(p.amount) if p else None
    if entity_type == "bank_statement_line":
        ln = db.get(BankStatementLine, entity_id)
        return abs(Decimal(ln.amount)) if ln else None
    if entity_type == "fixed_asset":
        a = db.get(FixedAsset, entity_id)
        return Decimal(a.purchase_cost + a.vat_amount) if a else None
    if entity_type == "stock_movement":
        mv = db.get(StockMovement, entity_id)
        return abs(Decimal(mv.total_cost)) if mv else None
    if entity_type == "payroll_run":
        r = db.get(PayrollRun, entity_id)
        return Decimal(r.net_total) if r else None
    if entity_type == "expense":
        ex = db.get(Expense, entity_id)
        return Decimal(ex.total_amount) if ex else None
    if entity_type == "customer_credit_note":
        cn = db.get(CustomerCreditNote, entity_id)
        return Decimal(cn.grand_total) if cn else None
    if entity_type == "vendor_credit_note":
        cn = db.get(VendorCreditNote, entity_id)
        return Decimal(cn.grand_total) if cn else None
    if entity_type == "project":
        p = db.get(Project, entity_id)
        return Decimal(p.contract_value) if p else None
    return None


_ENTITY_MODEL = {
    "journal_entry": JournalEntry, "sales_invoice": SalesInvoice, "vendor_bill": VendorBill,
    "customer_payment": CustomerPayment, "bill_payment": BillPayment,
    "bank_statement_line": BankStatementLine, "fixed_asset": FixedAsset,
    "stock_movement": StockMovement, "payroll_run": PayrollRun, "expense": Expense,
    "customer_credit_note": CustomerCreditNote, "vendor_credit_note": VendorCreditNote,
    "project": Project, "customer_advance": CustomerAdvance, "vendor_advance": VendorAdvance,
    "einvoice": EInvoice,
}


def _entity_exists(db: Session, entity_type: str, entity_id: str) -> bool:
    model = _ENTITY_MODEL.get(entity_type)
    if model is None:
        return entity_type in _ENTITY_TYPES  # linkable type without a strict FK check here
    return db.get(model, entity_id) is not None


# ── Audit trail ─────────────────────────────────────────────────────────────────────────────
def _log(db: Session, att: Attachment, action: str, actor: str, note: str | None = None) -> None:
    db.add(AttachmentEvent(attachment_id=att.id, action=action, actor=actor or "local", note=note))


# ── Document intelligence (template-free) ────────────────────────────────────────────────────
def _read_text(ref: str, ext: str) -> str | None:
    """Best-effort plain-text extraction from a stored file reference (filesystem path or Blob
    URL). Returns None when the type isn't machine-readable here (e.g. images need OCR)."""
    import io
    try:
        data = storage.read(ref)
    except Exception:  # noqa: BLE001 — storage/network hiccup shouldn't fail the upload
        return None
    try:
        if ext == "pdf":
            import pdfplumber
            with pdfplumber.open(io.BytesIO(data)) as pdf:
                return "\n".join((p.extract_text() or "") for p in pdf.pages)
        if ext in ("csv", "txt"):
            return data.decode("utf-8", errors="replace")
        if ext == "xlsx":
            from openpyxl import load_workbook
            wb = load_workbook(io.BytesIO(data), read_only=True, data_only=True)
            out = []
            for ws in wb.worksheets:
                for row in ws.iter_rows(values_only=True):
                    out.append(" ".join("" if c is None else str(c) for c in row))
            return "\n".join(out)
        if ext == "xls":
            import xlrd
            book = xlrd.open_workbook(file_contents=data)
            out = []
            for sh in book.sheets():
                for r in range(sh.nrows):
                    out.append(" ".join(str(sh.cell_value(r, c)) for c in range(sh.ncols)))
            return "\n".join(out)
        if ext == "docx":
            try:
                import docx  # python-docx, optional
            except Exception:  # noqa: BLE001
                return None
            d = docx.Document(io.BytesIO(data))
            parts = [p.text for p in d.paragraphs]
            for t in d.tables:
                for row in t.rows:
                    parts.append(" ".join(c.text for c in row.cells))
            return "\n".join(parts)
    except Exception:  # noqa: BLE001
        return None
    return None  # images, .doc (legacy) — not parsed in this prototype


_MONEY = r"(?:AED|USD|EUR|GBP|SAR|Dhs?\.?)?\s*((?:\d{1,3}(?:,\d{3})+|\d+)(?:\.\d{1,2})?)"
_CURR = re.compile(r"\b(AED|USD|EUR|GBP|SAR)\b", re.I)
_TRN = re.compile(r"\b(\d{15})\b")
_DATE = re.compile(r"\b(\d{4}-\d{2}-\d{2}|\d{1,2}[/\-.](?:\d{1,2}|[A-Za-z]{3,9})[/\-.]\d{2,4})\b")


def _to_dec(s: str) -> Decimal | None:
    try:
        return Decimal(s.replace(",", "").strip())
    except (InvalidOperation, AttributeError):
        return None


def _labeled_amount(text: str, labels: list[str]) -> Decimal | None:
    for lab in labels:
        m = re.search(lab + r"[^\d\-]{0,20}?" + _MONEY, text, re.I)
        if m:
            d = _to_dec(m.group(1))
            if d is not None:
                return d
    return None


def _parse_fields(text: str) -> dict:
    """Heuristic field extraction from arbitrary document text — no fixed template."""
    fields: dict = {}
    if not text:
        return fields
    flat = re.sub(r"[ \t]+", " ", text)

    # \b after the keyword stops 'inv' matching the prefix of 'Invoice' (which previously
    # captured the leftover 'oice'); separators include comma/newline/dash/= for CSV & label
    # layouts; the digit-guard rejects captured words that aren't real document numbers.
    m = re.search(r"\b(?:tax\s+invoice|invoice|inv|bill|credit\s+note)\b"
                  r"\s*(?:no\.?|number|nbr|ref(?:erence)?|#)?[\s:#.,\-=|]*"
                  r"([A-Za-z0-9][A-Za-z0-9\-/]{2,})", flat, re.I)
    if m and any(ch.isdigit() for ch in m.group(1)):
        fields["invoice_number"] = m.group(1).strip(" .:-,")
    m = re.search(r"(?:P\.?O\.?|purchase\s+order|order)\s*(?:no\.?|number|#|:)?[\s:#.,\-=|]*"
                  r"([A-Za-z0-9][A-Za-z0-9\-/]{2,})", flat, re.I)
    if m and any(ch.isdigit() for ch in m.group(1)):
        fields["order_reference"] = m.group(1).strip(" .:-,")
    trn = _TRN.search(flat)
    if trn:
        fields["trn"] = trn.group(1)
    dt = _DATE.search(flat)
    if dt:
        fields["date"] = dt.group(1)
    cur = _CURR.search(flat)
    if cur:
        fields["currency"] = cur.group(1).upper()

    net = _labeled_amount(flat, [r"net(?:\s+amount)?", r"sub[\s-]?total", r"amount\s+excl"])
    vat = _labeled_amount(flat, [r"vat(?:\s+amount)?", r"tax(?:\s+amount)?"])
    gross = _labeled_amount(flat, [r"grand\s*total", r"total\s+amount", r"amount\s+due",
                                   r"balance\s+due", r"total\s+incl", r"total"])
    if gross is None:
        # fall back to the largest money-looking number in the document
        nums = [d for d in (_to_dec(x) for x in re.findall(_MONEY, flat)) if d is not None]
        gross = max(nums) if nums else None
    if net is not None:
        fields["net_amount"] = str(net)
    if vat is not None:
        fields["vat_amount"] = str(vat)
    if gross is not None:
        fields["gross_amount"] = str(gross)
    return fields


def extract_and_match(db: Session, att: Attachment) -> None:
    """Run extraction on the stored file and compare its total against the transaction."""
    if att.file_ext in _IMAGE_EXT or att.file_ext == "doc":
        att.extraction_status = "unsupported"   # OCR / legacy .doc not available in this prototype
        att.extracted_json = json.dumps({"note": "Extraction not available for this file type "
                                                  "(image OCR / legacy .doc). Manual review required."})
        att.match_status = "unknown"
        return
    text = _read_text(att.storage_path, att.file_ext)
    if text is None:
        att.extraction_status = "failed"
        att.match_status = "unknown"
        return
    fields = _parse_fields(text)
    att.extracted_json = json.dumps(fields)
    att.extraction_status = "done"

    expected = _entity_amount(db, att.entity_type, att.entity_id)
    got = _to_dec(fields.get("gross_amount", "")) if fields.get("gross_amount") else None
    if expected is None or got is None:
        att.match_status = "unknown"
        att.match_difference = None
    else:
        diff = (Decimal(expected) - got).quantize(Decimal("0.01"))
        att.match_difference = diff
        att.match_status = "matched" if abs(diff) < Decimal("0.01") else "mismatch"


# ── Storage helpers ───────────────────────────────────────────────────────────────────────
def _scan_clean(data: bytes) -> bool:
    """Malware-scan hook. No scanner is bundled in this prototype, so this is a permissive
    stub — wire an AV engine (e.g. ClamAV `clamd`) here for production."""
    return True


def _validate(filename: str, data: bytes) -> str:
    ext = _ext_of(filename)
    if ext not in ALLOWED_EXT:
        raise AttachmentError(f"File type '.{ext or '?'}' is not allowed. Allowed: {', '.join(sorted(ALLOWED_EXT))}.")
    max_bytes = get_settings().max_upload_mb * 1024 * 1024
    if len(data) == 0:
        raise AttachmentError("The uploaded file is empty.")
    if len(data) > max_bytes:
        raise AttachmentError(f"File exceeds the {get_settings().max_upload_mb} MB limit.")
    if not _scan_clean(data):
        raise AttachmentError("File failed the malware scan.")
    return ext


def _write(att_id: str, entity_type: str, entity_id: str, ext: str, data: bytes,
           mime: str = "application/octet-stream") -> str:
    """Persist the bytes via the durable storage backend (Vercel Blob in production, filesystem
    locally / on a mounted disk). Returns the reference to store in the DB."""
    key = f"{entity_type}/{entity_id}/{att_id}.{ext}"
    return storage.save(key, data, mime)


def read_bytes(att: Attachment) -> bytes:
    """Fetch an attachment's bytes from durable storage (path or Blob URL)."""
    return storage.read(att.storage_path)


# ── Public API ────────────────────────────────────────────────────────────────────────────
def save_upload(db: Session, entity_type: str, entity_id: str, filename: str,
                data: bytes, actor: str = "local") -> Attachment:
    if entity_type not in _ENTITY_TYPES:
        raise AttachmentError(f"Unknown transaction type '{entity_type}'.")
    if not _entity_exists(db, entity_type, entity_id):
        raise AttachmentError("The transaction to attach to was not found.")
    ext = _validate(filename, data)
    att = Attachment(
        entity_type=entity_type, entity_id=str(entity_id),
        display_name=filename, original_name=filename, file_ext=ext,
        mime_type=_MIME.get(ext, "application/octet-stream"), file_size=len(data),
        storage_path="", sha256=hashlib.sha256(data).hexdigest(), uploaded_by=actor or "local",
    )
    db.add(att)
    db.flush()  # need att.id for the storage key
    att.storage_path = _write(att.id, entity_type, str(entity_id), ext, data, att.mime_type)
    _log(db, att, "uploaded", actor, note=f"{filename} ({len(data)} bytes)")
    extract_and_match(db, att)
    _log(db, att, "extracted", actor, note=f"extraction={att.extraction_status}, match={att.match_status}")
    db.commit()
    db.refresh(att)
    return att


def list_for(db: Session, entity_type: str, entity_id: str, include_deleted: bool = False) -> list[Attachment]:
    q = select(Attachment).where(Attachment.entity_type == entity_type,
                                 Attachment.entity_id == str(entity_id))
    if not include_deleted:
        q = q.where(Attachment.is_deleted.is_(False))
    return list(db.execute(q.order_by(Attachment.uploaded_at)).scalars())


def get(db: Session, att_id: str, allow_deleted: bool = False) -> Attachment:
    att = db.get(Attachment, att_id)
    if not att or (att.is_deleted and not allow_deleted):
        raise AttachmentError("Attachment not found.")
    return att


def rename(db: Session, att_id: str, new_name: str, actor: str = "local") -> Attachment:
    att = get(db, att_id)
    new_name = (new_name or "").strip()
    if not new_name:
        raise AttachmentError("A name is required.")
    old = att.display_name
    att.display_name = new_name
    _log(db, att, "renamed", actor, note=f"{old} → {new_name}")
    db.commit()
    db.refresh(att)
    return att


def replace(db: Session, att_id: str, filename: str, data: bytes, actor: str = "local") -> Attachment:
    att = get(db, att_id)
    ext = _validate(filename, data)
    # write the new file (new extension may differ), remove the old file if the path changes
    old_path = att.storage_path
    att.file_ext = ext
    att.mime_type = _MIME.get(ext, "application/octet-stream")
    att.original_name = filename
    att.file_size = len(data)
    att.sha256 = hashlib.sha256(data).hexdigest()
    att.storage_path = _write(att.id, att.entity_type, att.entity_id, ext, data, att.mime_type)
    if old_path and old_path != att.storage_path:
        storage.delete(old_path)
    att.review_status = "pending"
    extract_and_match(db, att)
    _log(db, att, "replaced", actor, note=f"→ {filename} ({len(data)} bytes)")
    db.commit()
    db.refresh(att)
    return att


def mark_reviewed(db: Session, att_id: str, actor: str = "local", note: str | None = None) -> Attachment:
    att = get(db, att_id)
    att.review_status = "reviewed"
    _log(db, att, "reviewed", actor, note=note)
    db.commit()
    db.refresh(att)
    return att


def reextract(db: Session, att_id: str, actor: str = "local") -> Attachment:
    att = get(db, att_id)
    extract_and_match(db, att)
    _log(db, att, "extracted", actor, note=f"re-run: extraction={att.extraction_status}, match={att.match_status}")
    db.commit()
    db.refresh(att)
    return att


def soft_delete(db: Session, att_id: str, actor: str = "local") -> None:
    att = get(db, att_id)
    att.is_deleted = True
    att.deleted_at = _now()
    att.deleted_by = actor or "local"
    _log(db, att, "deleted", actor)
    db.commit()


def record_access(db: Session, att_id: str, actor: str = "local", download: bool = False) -> Attachment:
    att = get(db, att_id)
    _log(db, att, "downloaded" if download else "viewed", actor)
    db.commit()
    db.refresh(att)
    return att


_PREVIEWABLE = {"pdf", "jpg", "jpeg", "png", "gif", "webp", "txt", "csv"}


def serialize(db: Session, att: Attachment) -> dict:
    try:
        extracted = json.loads(att.extracted_json) if att.extracted_json else {}
    except (ValueError, TypeError):
        extracted = {}
    return {
        "id": att.id, "entity_type": att.entity_type, "entity_id": att.entity_id,
        "display_name": att.display_name, "original_name": att.original_name,
        "file_ext": att.file_ext, "mime_type": att.mime_type, "file_size": att.file_size,
        "sha256": att.sha256, "uploaded_by": att.uploaded_by,
        "uploaded_at": _iso(att.uploaded_at), "modified_at": _iso(att.modified_at),
        "review_status": att.review_status, "extraction_status": att.extraction_status,
        "match_status": att.match_status, "match_difference": att.match_difference,
        "extracted": extracted,
        "transaction_amount": _entity_amount(db, att.entity_type, att.entity_id),
        "can_preview": att.file_ext in _PREVIEWABLE,
    }


def events_for(db: Session, att_id: str) -> list[AttachmentEvent]:
    att = get(db, att_id, allow_deleted=True)
    return list(att.events)


# ── Transaction-level status (for badges / filters) ─────────────────────────────────────────
def status_for(db: Session, entity_type: str, entity_id: str) -> dict:
    atts = list_for(db, entity_type, entity_id)
    n = len(atts)
    if n == 0:
        code, label = "none", "No document"
    elif any(a.match_status == "mismatch" for a in atts):
        code, label = "mismatch", "Document mismatch"
    elif any(a.review_status == "pending" for a in atts):
        code, label = "pending", "Pending review"
    elif all(a.review_status == "reviewed" for a in atts):
        code, label = "reviewed", "Reviewed"
    else:
        code, label = "attached", "Document attached"
    return {"code": code, "label": label, "count": n,
            "mismatch": sum(1 for a in atts if a.match_status == "mismatch"),
            "pending": sum(1 for a in atts if a.review_status == "pending")}


def status_bulk(db: Session, entity_type: str, ids: list[str]) -> dict[str, dict]:
    """Attachment status for many transactions at once — powers list-view badges & filters."""
    return {str(i): status_for(db, entity_type, str(i)) for i in ids}
